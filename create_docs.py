from docx import Document
from docx.shared import Inches

from docx.enum.style import WD_STYLE_TYPE
from docx import *
document = Document()
styles = document.styles

#生成所有表样式
for s in styles:
    if s.type == WD_STYLE_TYPE.TABLE:
        print(s)





















document = Document()

document.add_heading('Document Title', 0)  #插入标题

p = document.add_paragraph('A plain paragraph having some ')   #插入段落
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

# document.add_picture('monty-truth.png', width=Inches(1.25)) #插入图片

table = document.add_table(rows=1, cols=3) #插入表格
table.style ='Medium Grid 1 Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
recordset=[('da','fgr',10)]
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item[0])
    row_cells[1].text = str(item[1])
    rt=item[2]
    row_cells[2].text = str(rt)

document.add_page_break()

document.save('demo.docx')  #保存文档